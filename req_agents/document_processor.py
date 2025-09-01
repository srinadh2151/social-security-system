"""
Multimodal Document Processing Agent

This agent processes various document types (PDFs, Word docs, Excel files) 
and extracts structured information for financial assessment applications.
"""

import logging
import json
import os
import io
import base64
from typing import Dict, Any, List, Optional, Union, Tuple
from datetime import datetime
from enum import Enum
import asyncio
from pathlib import Path

# Document processing libraries
import PyPDF2
import pdfplumber
from PIL import Image
import pandas as pd
from docx import Document
import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows

# OCR and image processing
try:
    import pytesseract
    from pdf2image import convert_from_path, convert_from_bytes
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR libraries not available. Install pytesseract and pdf2image for full functionality.")

# Vision capabilities
try:
    import cv2
    import numpy as np
    VISION_AVAILABLE = True
except ImportError:
    VISION_AVAILABLE = False
    logging.warning("Computer vision libraries not available. Install opencv-python for enhanced image processing.")

try:
    from req_agents.llm_interface import LangChainLLMInterface
    from req_agents.base_agent import LangChainBaseAgent
except ImportError:
    from llm_interface import LangChainLLMInterface
    from base_agent import LangChainBaseAgent

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types."""
    EMIRATES_ID = "emirates_id"
    RESUME = "resume"
    ASSETS_LIABILITIES = "assets_liabilities"
    CREDIT_REPORT = "credit_report"
    BANK_STATEMENT = "bank_statement"
    UNKNOWN = "unknown"


class FileFormat(Enum):
    """Supported file formats."""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    XLSX = "xlsx"
    XLS = "xls"
    IMAGE = "image"
    TXT = "txt"


class DocumentProcessingAgent(LangChainBaseAgent):
    """Agent for processing multimodal documents."""
    
    def __init__(
        self,
        llm_interface: Optional[LangChainLLMInterface] = None,
        model: str = "gpt-4o",  # Use GPT-4o for vision capabilities
        temp_dir: str = "./temp_docs"
    ):
        """Initialize document processing agent."""
        super().__init__(
            name="document_processor",
            llm_interface=llm_interface,
            model=model,
            system_prompt=self._get_system_prompt()
        )
        
        self.temp_dir = Path(temp_dir)
        self.temp_dir.mkdir(exist_ok=True)
        
        # Document type patterns for classification
        self.document_patterns = {
            DocumentType.EMIRATES_ID: [
                "emirates id", "identity card", "uae id", "national id",
                "هوية الإمارات", "بطاقة الهوية", "resident flag", "nationality"
            ],
            DocumentType.RESUME: [
                "resume", "cv", "curriculum vitae", "experience",
                "education", "skills", "employment history", "work experience",
                "professional experience", "career summary"
            ],
            DocumentType.ASSETS_LIABILITIES: [
                "assets", "liabilities", "balance sheet", "financial statement",
                "net worth", "portfolio", "investments", "bank account",
                "savings", "property", "loans", "credit cards"
            ],
            DocumentType.CREDIT_REPORT: [
                "credit report", "credit score", "credit history", "credit bureau",
                "aecb", "al etihad credit bureau", "etihad bureau", "experian", "equifax",
                "cb subject id", "provider no", "credit summary", "payment history",
                "information providers", "response id"
            ],
            DocumentType.BANK_STATEMENT: [
                "bank statement", "account statement", "transaction history", "account details",
                "opening balance", "closing balance", "debit", "credit", "account balance",
                "emirates nbd", "adcb", "fab", "mashreq", "rakbank", "statement from"
            ]
        }
    
    def _get_system_prompt(self) -> str:
        """Get system prompt for document processing."""
        return """
You are an expert document processing agent specializing in financial and identity documents.
Your role is to extract structured information from various document types including:

1. Emirates ID documents - Extract personal information, demographics
2. Resumes/CVs - Extract employment history, education, skills
3. Assets/Liabilities statements - Extract financial data, net worth
4. Credit reports - Extract credit history, scores, financial obligations

Key principles:
1. Be accurate and thorough in data extraction
2. Maintain data privacy and security
3. Structure information consistently
4. Handle multiple languages (English/Arabic)
5. Validate extracted information for consistency
6. Flag any suspicious or unclear information

Always provide structured JSON output with confidence scores for extracted data.
"""
    
    async def execute(self, task: Dict[str, Any]) -> Dict[str, Any]:
        """Execute document processing task."""
        try:
            file_path = task.get("file_path")
            document_purpose = task.get("document_purpose", "unknown")
            
            if not file_path or not os.path.exists(file_path):
                return {
                    "status": "error",
                    "message": "File not found or path not provided",
                    "extracted_data": {}
                }
            
            # Process the document
            result = await self.process_document(file_path, document_purpose)
            
            return {
                "status": "success",
                "file_path": file_path,
                "document_purpose": document_purpose,
                "extracted_data": result,
                "processing_timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Document processing failed: {str(e)}")
            return {
                "status": "error",
                "message": str(e),
                "extracted_data": {}
            }
    
    async def process_document(
        self,
        file_path: str,
        document_purpose: str = "unknown"
    ) -> Dict[str, Any]:
        """Process a document and extract structured information."""
        file_path = Path(file_path)
        
        # Determine file format
        file_format = self._get_file_format(file_path)
        
        # Extract content based on file format first
        extracted_content = await self._extract_content(file_path, file_format)
        
        # Classify document type (now with access to extracted content if needed)
        document_type = await self._classify_document_type(file_path, document_purpose, extracted_content)
        logger.info(f"Classified {file_path.name} as {document_type.value} (purpose: {document_purpose})")
        
        # Process content based on document type
        structured_data = await self._process_by_document_type(
            extracted_content, document_type, file_format
        )
        
        return {
            "document_type": document_type.value,
            "file_format": file_format.value,
            "extracted_content": extracted_content,
            "structured_data": structured_data,
            "confidence_score": structured_data.get("confidence_score", 0.0)
        }
    
    def _get_file_format(self, file_path: Path) -> FileFormat:
        """Determine file format from extension."""
        extension = file_path.suffix.lower()
        
        format_mapping = {
            ".pdf": FileFormat.PDF,
            ".docx": FileFormat.DOCX,
            ".doc": FileFormat.DOC,
            ".txt": FileFormat.TXT,
            ".xlsx": FileFormat.XLSX,
            ".xls": FileFormat.XLS,
            ".png": FileFormat.IMAGE,
            ".jpg": FileFormat.IMAGE,
            ".jpeg": FileFormat.IMAGE,
            ".tiff": FileFormat.IMAGE,
            ".bmp": FileFormat.IMAGE
        }
        
        return format_mapping.get(extension, FileFormat.PDF)
    
    async def _classify_document_type(
        self,
        file_path: Path,
        document_purpose: str,
        extracted_content: Optional[Dict[str, Any]] = None
    ) -> DocumentType:
        """Classify document type based on purpose and content."""
        # First, try to classify based on provided purpose
        purpose_lower = document_purpose.lower()
        logger.info(f"Classifying document with purpose: '{document_purpose}' (lowercase: '{purpose_lower}')")
        
        for doc_type, patterns in self.document_patterns.items():
            matching_patterns = [pattern for pattern in patterns if pattern in purpose_lower]
            if matching_patterns:
                logger.info(f"Matched {doc_type.value} based on purpose patterns: {matching_patterns}")
                return doc_type
        
        # Special case: if purpose exactly matches a document type, use it
        purpose_to_type = {
            "emirates_id": DocumentType.EMIRATES_ID,
            "credit_report": DocumentType.CREDIT_REPORT,
            "resume": DocumentType.RESUME,
            "assets_liabilities": DocumentType.ASSETS_LIABILITIES,
            "bank_statement": DocumentType.BANK_STATEMENT
        }
        
        if purpose_lower in purpose_to_type:
            logger.info(f"Matched {purpose_to_type[purpose_lower].value} based on exact purpose match")
            return purpose_to_type[purpose_lower]
        
        # If purpose doesn't match, try to classify from filename
        filename_lower = file_path.name.lower()
        for doc_type, patterns in self.document_patterns.items():
            if any(pattern in filename_lower for pattern in patterns):
                return doc_type
        
        # If still unknown, try to classify based on extracted content
        if extracted_content:
            try:
                # Get text content for analysis
                sample_content = extracted_content.get("text", "").lower()
                
                # Check content against patterns with scoring system
                if sample_content:
                    # Score each document type based on indicators found
                    type_scores = {}
                    
                    # Credit report specific indicators (high confidence)
                    credit_indicators = [
                        "credit report", "credit score", "credit bureau", "aecb", 
                        "al etihad credit bureau", "etihad bureau", "cb subject id",
                        "provider no", "response id", "information providers"
                    ]
                    
                    credit_matches = [indicator for indicator in credit_indicators if indicator in sample_content]
                    if credit_matches:
                        type_scores[DocumentType.CREDIT_REPORT] = len(credit_matches) * 2  # Higher weight
                        logger.info(f"Credit report indicators found: {credit_matches}")
                    
                    # Emirates ID specific indicators (high confidence)
                    emirates_id_indicators = [
                        "emirates id", "identity card", "uae id", "expiry date", "resident card",
                        "issue date", "place of birth"
                    ]
                    
                    emirates_matches = [indicator for indicator in emirates_id_indicators if indicator in sample_content]
                    if emirates_matches:
                        type_scores[DocumentType.EMIRATES_ID] = len(emirates_matches) * 2  # Higher weight
                        logger.info(f"Emirates ID indicators found: {emirates_matches}")
                    
                    # Resume specific indicators (need multiple matches)
                    resume_indicators = [
                        "work experience", "employment history", "professional experience",
                        "education", "skills", "qualifications", "certifications"
                    ]
                    
                    resume_matches = [indicator for indicator in resume_indicators if indicator in sample_content]
                    if len(resume_matches) >= 2:  # Need at least 2 for resume
                        type_scores[DocumentType.RESUME] = len(resume_matches)
                        logger.info(f"Resume indicators found: {resume_matches}")
                    
                    # Assets/Liabilities indicators
                    assets_indicators = [
                        "balance sheet", "assets", "liabilities", "net worth",
                        "investments", "bank account", "financial statement"
                    ]
                    
                    assets_matches = [indicator for indicator in assets_indicators if indicator in sample_content]
                    if len(assets_matches) >= 2:  # Need multiple for assets
                        type_scores[DocumentType.ASSETS_LIABILITIES] = len(assets_matches)
                        logger.info(f"Assets/Liabilities indicators found: {assets_matches}")
                    
                    # Bank statement specific indicators (high confidence)
                    bank_statement_indicators = [
                        "account statement", "transaction", "debit", "credit", "account balance",
                        "opening balance", "closing balance", "statement from", "account number",
                        "emirates nbd", "adcb", "fab", "mashreq", "rakbank"
                    ]
                    
                    bank_matches = [indicator for indicator in bank_statement_indicators if indicator in sample_content]
                    if len(bank_matches) >= 3:  # Need multiple matches for bank statement
                        type_scores[DocumentType.BANK_STATEMENT] = len(bank_matches) * 2  # Higher weight
                        logger.info(f"Bank statement indicators found: {bank_matches}")
                    
                    # Return the document type with highest score
                    if type_scores:
                        best_type = max(type_scores, key=type_scores.get)
                        best_score = type_scores[best_type]
                        logger.info(f"Classified as {best_type.value} with score {best_score}")
                        return best_type
                    
                    # Fallback: check general patterns but with lower confidence
                    for doc_type, patterns in self.document_patterns.items():
                        pattern_matches = [pattern for pattern in patterns if pattern in sample_content]
                        if len(pattern_matches) >= 1:  # At least one match for fallback
                            logger.info(f"Fallback classification as {doc_type.value} based on patterns: {pattern_matches}")
                            return doc_type
            
            except Exception as e:
                logger.warning(f"Content-based classification failed: {str(e)}")
        
        return DocumentType.UNKNOWN
    
    async def _extract_content(
        self,
        file_path: Path,
        file_format: FileFormat
    ) -> Dict[str, Any]:
        """Extract content from file based on format."""
        content = {
            "text": "",
            "tables": [],
            "images": [],
            "metadata": {}
        }
        
        try:
            if file_format == FileFormat.PDF:
                content = await self._extract_pdf_content(file_path)
            elif file_format == FileFormat.DOCX:
                content = await self._extract_docx_content(file_path)
            elif file_format in [FileFormat.XLSX, FileFormat.XLS]:
                content = await self._extract_excel_content(file_path)
            elif file_format == FileFormat.IMAGE:
                content = await self._extract_image_content(file_path)
            elif file_format == FileFormat.TXT:
                content = await self._extract_txt_content(file_path)
            
        except Exception as e:
            logger.error(f"Content extraction failed for {file_path}: {str(e)}")
            content["error"] = str(e)
        
        return content
    
    async def _extract_pdf_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from PDF files."""
        content = {
            "text": "",
            "tables": [],
            "images": [],
            "metadata": {}
        }
        
        # Try pdfplumber first (better for tables and layout)
        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables):
                        if table:
                            content["tables"].append({
                                "page": page_num + 1,
                                "table_number": table_num + 1,
                                "data": table,
                                "headers": table[0] if table else []
                            })
                
                content["text"] = "\n\n".join(text_parts)
                content["metadata"] = pdf.metadata or {}
                
                logger.info(f"Successfully extracted {len(content['text'])} characters using pdfplumber")
                if content["text"]:
                    logger.info(f"Sample text: {content['text'][:200]}...")
        
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {str(e)}, trying PyPDF2 fallback")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_parts = []
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                        except Exception as page_error:
                            logger.warning(f"Failed to extract text from page {page_num + 1}: {str(page_error)}")
                    
                    content["text"] = "\n\n".join(text_parts)
                    content["metadata"] = pdf_reader.metadata or {}
                    
                    logger.info(f"Successfully extracted {len(content['text'])} characters using PyPDF2 fallback")
                    if content["text"]:
                        logger.info(f"Sample text: {content['text'][:200]}...")
            
            except Exception as fallback_error:
                logger.error(f"Both PDF extraction methods failed: {str(fallback_error)}")
                content["error"] = f"PDF extraction failed: {str(fallback_error)}"
        
        # Try to extract images if OCR is available (only if poppler is available)
        if OCR_AVAILABLE and content.get("text"):  # Only try if we have some text already
            try:
                images = convert_from_path(str(file_path))
                for i, image in enumerate(images):
                    # Save image temporarily
                    temp_image_path = self.temp_dir / f"temp_page_{i}.png"
                    image.save(temp_image_path)
                    
                    # Extract text from image using OCR
                    ocr_text = pytesseract.image_to_string(image)
                    
                    # Convert image to base64 for LLM processing
                    img_buffer = io.BytesIO()
                    image.save(img_buffer, format='PNG')
                    img_base64 = base64.b64encode(img_buffer.getvalue()).decode()
                    
                    content["images"].append({
                        "page": i + 1,
                        "ocr_text": ocr_text,
                        "base64": img_base64,
                        "temp_path": str(temp_image_path)
                    })
                    
            except Exception as e:
                logger.warning(f"Image extraction failed: {str(e)}")
        
        return content
    
    async def _extract_docx_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from Word documents."""
        content = {
            "text": "",
            "tables": [],
            "images": [],
            "metadata": {}
        }
        
        try:
            doc = Document(file_path)
            
            # Extract text
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            content["text"] = "\n".join(text_parts)
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                content["tables"].append({
                    "table_number": table_num + 1,
                    "data": table_data,
                    "headers": table_data[0] if table_data else []
                })
            
            # Extract metadata
            content["metadata"] = {
                "author": doc.core_properties.author,
                "title": doc.core_properties.title,
                "created": str(doc.core_properties.created),
                "modified": str(doc.core_properties.modified)
            }
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {str(e)}")
            content["error"] = str(e)
        
        return content
    
    async def _extract_excel_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from Excel files."""
        content = {
            "text": "",
            "tables": [],
            "images": [],
            "metadata": {}
        }
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert DataFrame to table format
                table_data = []
                
                # Add headers
                headers = df.columns.tolist()
                table_data.append(headers)
                
                # Add data rows
                for _, row in df.iterrows():
                    table_data.append(row.tolist())
                
                content["tables"].append({
                    "sheet_name": sheet_name,
                    "data": table_data,
                    "headers": headers,
                    "shape": df.shape
                })
                
                # Create text summary
                text_summary = f"Sheet: {sheet_name}\n"
                text_summary += f"Dimensions: {df.shape[0]} rows x {df.shape[1]} columns\n"
                text_summary += f"Columns: {', '.join(headers)}\n"
                
                content["text"] += text_summary + "\n"
            
            # Extract metadata
            workbook = openpyxl.load_workbook(file_path)
            content["metadata"] = {
                "sheet_names": excel_file.sheet_names,
                "total_sheets": len(excel_file.sheet_names)
            }
            
        except Exception as e:
            logger.error(f"Excel processing failed: {str(e)}")
            content["error"] = str(e)
        
        return content
    
    async def _extract_image_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from image files."""
        content = {
            "text": "",
            "tables": [],
            "images": [],
            "metadata": {}
        }
        
        try:
            # Load image
            image = Image.open(file_path)
            
            # Convert to base64
            img_buffer = io.BytesIO()
            image.save(img_buffer, format='PNG')
            img_base64 = base64.b64encode(img_buffer.getvalue()).decode()
            
            # Extract text using OCR if available
            ocr_text = ""
            if OCR_AVAILABLE:
                ocr_text = pytesseract.image_to_string(image)
            
            content["images"].append({
                "base64": img_base64,
                "ocr_text": ocr_text,
                "size": image.size,
                "format": image.format
            })
            
            content["text"] = ocr_text
            content["metadata"] = {
                "size": image.size,
                "format": image.format,
                "mode": image.mode
            }
            
        except Exception as e:
            logger.error(f"Image processing failed: {str(e)}")
            content["error"] = str(e)
        
        return content
    
    async def _extract_txt_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from text files."""
        content = {
            "text": "",
            "tables": [],
            "images": [],
            "metadata": {}
        }
        
        try:
            # Read text file with different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text_content = ""
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text_content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not text_content:
                # Fallback: read as binary and decode with errors='ignore'
                with open(file_path, 'rb') as f:
                    text_content = f.read().decode('utf-8', errors='ignore')
            
            content["text"] = text_content
            
            # Get file metadata
            file_stats = file_path.stat()
            content["metadata"] = {
                "file_size": file_stats.st_size,
                "created": datetime.fromtimestamp(file_stats.st_ctime).isoformat(),
                "modified": datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
                "encoding": "utf-8"  # Assume utf-8 after processing
            }
            
            # Try to detect table-like structures in text
            lines = text_content.split('\n')
            potential_tables = []
            
            for i, line in enumerate(lines):
                # Look for lines with multiple columns (separated by spaces, tabs, or |)
                if any(sep in line for sep in ['\t', '  ', '|']) and len(line.strip()) > 0:
                    # Check if this could be part of a table
                    columns = [col.strip() for col in line.split() if col.strip()]
                    if len(columns) > 2:  # At least 3 columns to consider it tabular
                        potential_tables.append({
                            "line_number": i + 1,
                            "content": line.strip(),
                            "columns": columns
                        })
            
            if potential_tables:
                content["tables"].append({
                    "table_number": 1,
                    "data": potential_tables,
                    "headers": [],
                    "note": "Detected potential tabular data in text"
                })
            
        except Exception as e:
            logger.error(f"Text file processing failed: {str(e)}")
            content["error"] = str(e)
        
        return content
    
    async def _process_by_document_type(
        self,
        extracted_content: Dict[str, Any],
        document_type: DocumentType,
        file_format: FileFormat
    ) -> Dict[str, Any]:
        """Process extracted content based on document type."""
        
        if document_type == DocumentType.EMIRATES_ID:
            return await self._process_emirates_id(extracted_content)
        elif document_type == DocumentType.RESUME:
            return await self._process_resume(extracted_content)
        elif document_type == DocumentType.ASSETS_LIABILITIES:
            return await self._process_assets_liabilities(extracted_content)
        elif document_type == DocumentType.CREDIT_REPORT:
            return await self._process_credit_report(extracted_content)
        elif document_type == DocumentType.BANK_STATEMENT:
            return await self._process_bank_statement(extracted_content)
        else:
            return await self._process_unknown_document(extracted_content)
    
    async def _process_emirates_id(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Process Emirates ID document."""
        prompt = f"""
        Extract structured information from this Emirates ID document:
        
        Text Content: {content.get('text', '')}
        
        OCR from Images: {json.dumps([img.get('ocr_text', '') for img in content.get('images', [])], indent=2)}
        
        Extract the following information and return as JSON:
        {{
            "personal_info": {{
                "full_name": "string",
                "full_name_arabic": "string", 
                "nationality": "string",
                "date_of_birth": "YYYY-MM-DD",
                "place_of_birth": "string",
                "gender": "male/female",
                "id_number": "string",
                "issue_date": "YYYY-MM-DD",
                "expiry_date": "YYYY-MM-DD"
            }},
            "address_info": {{
                "emirate": "string",
                "area": "string",
                "full_address": "string"
            }},
            "confidence_score": 0.0-1.0,
            "extracted_fields": ["list of successfully extracted fields"],
            "missing_fields": ["list of fields that couldn't be extracted"]
        }}
        
        Be very careful with dates and ID numbers. If information is unclear, mark confidence as low.
        """
        
        response = await self.generate_structured_response(
            prompt,
            schema={
                "type": "object",
                "properties": {
                    "personal_info": {"type": "object"},
                    "address_info": {"type": "object"},
                    "confidence_score": {"type": "number"},
                    "extracted_fields": {"type": "array"},
                    "missing_fields": {"type": "array"}
                }
            }
        )
        
        return response
    
    async def _process_resume(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Process resume/CV document."""
        prompt = f"""
        Extract structured information from this resume/CV:
        
        Text Content: {content.get('text', '')}
        
        Tables: {json.dumps(content.get('tables', []), indent=2)}
        
        Extract the following information and return as JSON:
        {{
            "personal_info": {{
                "name": "string",
                "email": "string",
                "phone": "string",
                "address": "string",
                "linkedin": "string",
                "nationality": "string"
            }},
            "employment_history": [
                {{
                    "company": "string",
                    "position": "string",
                    "start_date": "YYYY-MM or YYYY",
                    "end_date": "YYYY-MM or YYYY or 'present'",
                    "duration_months": "number",
                    "description": "string",
                    "employment_type": "full_time/part_time/contract/freelance"
                }}
            ],
            "education": [
                {{
                    "institution": "string",
                    "degree": "string",
                    "field_of_study": "string",
                    "graduation_year": "YYYY",
                    "gpa": "string"
                }}
            ],
            "skills": ["list of skills"],
            "certifications": ["list of certifications"],
            "languages": ["list of languages"],
            "total_experience_years": "number",
            "current_employment_status": "employed/unemployed/student",
            "confidence_score": 0.0-1.0
        }}
        
        Calculate total experience carefully and determine current employment status.
        """
        
        response = await self.generate_structured_response(
            prompt,
            schema={
                "type": "object",
                "properties": {
                    "personal_info": {"type": "object"},
                    "employment_history": {"type": "array"},
                    "education": {"type": "array"},
                    "skills": {"type": "array"},
                    "total_experience_years": {"type": "number"},
                    "confidence_score": {"type": "number"}
                }
            }
        )
        
        return response
    
    async def _process_assets_liabilities(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Process assets and liabilities statement."""
        prompt = f"""
        Extract structured financial information from this assets/liabilities statement:
        
        Tables: {json.dumps(content.get('tables', []), indent=2)}
        
        Text Content: {content.get('text', '')}
        
        Extract the following information and return as JSON:
        {{
            "assets": {{
                "cash_and_equivalents": "number",
                "bank_accounts": [
                    {{
                        "bank_name": "string",
                        "account_type": "string",
                        "balance": "number"
                    }}
                ],
                "investments": {{
                    "stocks": "number",
                    "bonds": "number",
                    "mutual_funds": "number",
                    "real_estate": "number",
                    "other_investments": "number"
                }},
                "property": [
                    {{
                        "type": "residential/commercial/land",
                        "location": "string",
                        "estimated_value": "number",
                        "mortgage_outstanding": "number"
                    }}
                ],
                "vehicles": [
                    {{
                        "type": "string",
                        "model": "string",
                        "year": "number",
                        "estimated_value": "number",
                        "loan_outstanding": "number"
                    }}
                ],
                "other_assets": "number",
                "total_assets": "number"
            }},
            "liabilities": {{
                "credit_cards": [
                    {{
                        "bank": "string",
                        "outstanding_balance": "number",
                        "credit_limit": "number"
                    }}
                ],
                "loans": [
                    {{
                        "type": "personal/auto/mortgage/business",
                        "lender": "string",
                        "outstanding_balance": "number",
                        "monthly_payment": "number"
                    }}
                ],
                "other_liabilities": "number",
                "total_liabilities": "number"
            }},
            "net_worth": "number",
            "statement_date": "YYYY-MM-DD",
            "confidence_score": 0.0-1.0
        }}
        
        Calculate totals carefully and ensure net worth = total assets - total liabilities.
        """
        
        response = await self.generate_structured_response(
            prompt,
            schema={
                "type": "object",
                "properties": {
                    "assets": {"type": "object"},
                    "liabilities": {"type": "object"},
                    "net_worth": {"type": "number"},
                    "confidence_score": {"type": "number"}
                }
            }
        )
        
        return response
    
    async def _process_credit_report(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Process credit report document."""
        prompt = f"""
        Extract structured information from this AECB (Al Etihad Credit Bureau) credit report:
        
        Text Content: {content.get('text', '')}
        
        Tables: {json.dumps(content.get('tables', []), indent=2)}
        
        OCR from Images: {json.dumps([img.get('ocr_text', '') for img in content.get('images', [])], indent=2)}
        
        This appears to be an AECB credit report. Extract the following information and return as JSON:
        {{
            "personal_info": {{
                "name": "string (full name)",
                "first_name": "string",
                "last_name": "string",
                "arabic_name": "string (if available)",
                "gender": "male/female",
                "date_of_birth": "YYYY-MM-DD",
                "nationality": "string",
                "cb_subject_id": "string",
                "emirates_id": "string",
                "passport_number": "string"
            }},
            "credit_score": {{
                "score": "number (if visible, otherwise null)",
                "score_range": "string (e.g., 300-850)",
                "rating": "excellent/good/fair/poor (if determinable)",
                "bureau": "AECB"
            }},
            "addresses": [
                {{
                    "address": "string",
                    "emirate": "string",
                    "po_box": "string",
                    "provider": "string",
                    "date_updated": "string"
                }}
            ],
            "contact_info": {{
                "email": "string",
                "phone_numbers": ["list of phone numbers"],
                "mobile_numbers": ["list of mobile numbers"]
            }},
            "information_providers": [
                {{
                    "provider_code": "string",
                    "provider_name": "string", 
                    "description": "string",
                    "last_update": "string"
                }}
            ],
            "identification": [
                {{
                    "type": "passport/emirates_id",
                    "number": "string",
                    "expiry_date": "string"
                }}
            ],
            "report_metadata": {{
                "report_date": "YYYY-MM-DD",
                "response_id": "string",
                "bureau": "AECB"
            }},
            "confidence_score": 0.0-1.0
        }}
        
        Pay special attention to:
        1. The credit score (visible as "300" in the summary section)
        2. Personal information in the Customer Information section
        3. Provider information and their update dates
        4. Contact details and addresses
        5. Identification documents (passport, Emirates ID)
        
        Be very careful with dates and ID numbers. Extract exactly as shown.
        """
        
        response = await self.generate_structured_response(
            prompt,
            schema={
                "type": "object",
                "properties": {
                    "personal_info": {"type": "object"},
                    "credit_score": {"type": "object"},
                    "addresses": {"type": "array"},
                    "contact_info": {"type": "object"},
                    "information_providers": {"type": "array"},
                    "identification": {"type": "array"},
                    "report_metadata": {"type": "object"},
                    "confidence_score": {"type": "number"}
                }
            }
        )
        
        return response
    
    async def _process_bank_statement(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Process bank statement document."""
        prompt = f"""
        Extract structured information from this bank statement:
        
        Text Content: {content.get('text', '')}
        
        Tables: {json.dumps(content.get('tables', []), indent=2)}
        
        This appears to be a bank account statement. Extract the following information and return as JSON:
        {{
            "account_info": {{
                "account_holder_name": "string",
                "account_number": "string",
                "iban": "string",
                "account_type": "string (e.g., CURRENT ACCOUNT, SAVINGS)",
                "bank_name": "string",
                "branch": "string",
                "currency": "string"
            }},
            "statement_period": {{
                "statement_from": "YYYY-MM-DD",
                "statement_to": "YYYY-MM-DD",
                "statement_date": "YYYY-MM-DD"
            }},
            "balance_summary": {{
                "opening_balance": "number",
                "closing_balance": "number",
                "average_balance": "number (if calculable)",
                "minimum_balance": "number",
                "maximum_balance": "number"
            }},
            "transaction_summary": {{
                "total_transactions": "number",
                "total_debits": "number",
                "total_credits": "number",
                "total_debit_amount": "number",
                "total_credit_amount": "number",
                "salary_credits": "number",
                "atm_withdrawals": "number",
                "service_charges": "number"
            }},
            "transactions": [
                {{
                    "date": "YYYY-MM-DD",
                    "description": "string",
                    "transaction_type": "debit/credit",
                    "amount": "number",
                    "balance": "number",
                    "category": "salary/withdrawal/transfer/service_charge/purchase/other"
                }}
            ],
            "income_analysis": {{
                "salary_frequency": "monthly/bi-weekly/weekly",
                "average_monthly_salary": "number",
                "other_income_sources": ["list of other income types"],
                "total_monthly_income": "number"
            }},
            "spending_analysis": {{
                "average_monthly_spending": "number",
                "atm_withdrawal_frequency": "number per month",
                "service_charges_monthly": "number",
                "largest_expense": "number",
                "spending_categories": {{
                    "atm_withdrawals": "number",
                    "transfers": "number", 
                    "service_charges": "number",
                    "purchases": "number",
                    "other": "number"
                }}
            }},
            "financial_behavior": {{
                "account_management": "good/fair/poor",
                "overdraft_incidents": "number",
                "bounce_incidents": "number",
                "average_daily_balance": "number",
                "cash_flow_pattern": "stable/irregular/declining/improving"
            }},
            "confidence_score": 0.0-1.0
        }}
        
        Pay special attention to:
        1. Extracting all transaction details accurately
        2. Calculating totals and averages correctly
        3. Identifying salary payments and their frequency
        4. Categorizing different types of transactions
        5. Analyzing spending patterns and financial behavior
        
        Be very careful with amounts and dates. Extract exactly as shown.
        """
        
        response = await self.generate_structured_response(
            prompt,
            schema={
                "type": "object",
                "properties": {
                    "account_info": {"type": "object"},
                    "statement_period": {"type": "object"},
                    "balance_summary": {"type": "object"},
                    "transaction_summary": {"type": "object"},
                    "transactions": {"type": "array"},
                    "income_analysis": {"type": "object"},
                    "spending_analysis": {"type": "object"},
                    "financial_behavior": {"type": "object"},
                    "confidence_score": {"type": "number"}
                }
            }
        )
        
        return response
    
    async def _process_unknown_document(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Process unknown document type."""
        prompt = f"""
        Analyze this document and extract any relevant financial or personal information:
        
        Text Content: {content.get('text', '')}
        
        Tables: {json.dumps(content.get('tables', []), indent=2)}
        
        Determine what type of document this might be and extract relevant information:
        {{
            "document_classification": "string",
            "extracted_info": {{
                "personal_details": {{}},
                "financial_details": {{}},
                "employment_details": {{}},
                "other_details": {{}}
            }},
            "confidence_score": 0.0-1.0,
            "recommendations": ["what additional processing might be needed"]
        }}
        """
        
        response = await self.generate_structured_response(
            prompt,
            schema={
                "type": "object",
                "properties": {
                    "document_classification": {"type": "string"},
                    "extracted_info": {"type": "object"},
                    "confidence_score": {"type": "number"},
                    "recommendations": {"type": "array"}
                }
            }
        )
        
        return response
    
    async def convert_to_assessment_format(
        self,
        processed_documents: List[Dict[str, Any]],
        application_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """Convert processed documents to assessment agent format."""
        
        # Use provided application_id or generate a fallback
        if not application_id:
            application_id = f"APP-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            logger.warning(f"No application_id provided, generated fallback: {application_id}")
        
        # Initialize assessment data structure
        assessment_data = {
            "application_id": application_id,
            "applicant_info": {},
            "income_info": {},
            "employment_info": {},
            "family_info": {},
            "wealth_info": {},
            "demographic_info": {},
            "document_sources": []
        }
        
        for doc in processed_documents:
            doc_type = doc.get("document_type")
            structured_data = doc.get("structured_data", {})
            
            # Track document sources
            assessment_data["document_sources"].append({
                "type": doc_type,
                "confidence": structured_data.get("confidence_score", 0.0),
                "file_format": doc.get("file_format")
            })
            
            # Map data based on document type
            if doc_type == "emirates_id":
                await self._map_emirates_id_data(structured_data, assessment_data)
            elif doc_type == "resume":
                await self._map_resume_data(structured_data, assessment_data)
            elif doc_type == "assets_liabilities":
                await self._map_assets_data(structured_data, assessment_data)
            elif doc_type == "credit_report":
                await self._map_credit_data(structured_data, assessment_data)
            elif doc_type == "bank_statement":
                await self._map_bank_statement_data(structured_data, assessment_data)
        
        return assessment_data
    
    async def _map_emirates_id_data(self, data: Dict[str, Any], assessment_data: Dict[str, Any]):
        """Map Emirates ID data to assessment format."""
        personal_info = data.get("personal_info", {})
        address_info = data.get("address_info", {})
        
        assessment_data["applicant_info"].update({
            "name": personal_info.get("full_name", ""),
            "age": self._calculate_age(personal_info.get("date_of_birth")),
            "address": address_info.get("full_address", ""),
            "nationality": personal_info.get("nationality", ""),
            "id_number": personal_info.get("id_number", "")
        })
        
        assessment_data["demographic_info"].update({
            "gender": personal_info.get("gender", ""),
            "nationality": personal_info.get("nationality", ""),
            "emirate": address_info.get("emirate", "")
        })
    
    async def _map_resume_data(self, data: Dict[str, Any], assessment_data: Dict[str, Any]):
        """Map resume data to assessment format."""
        personal_info = data.get("personal_info", {})
        employment_history = data.get("employment_history", [])
        education = data.get("education", [])
        
        # Update applicant info
        if not assessment_data["applicant_info"].get("name"):
            assessment_data["applicant_info"]["name"] = personal_info.get("name", "")
        
        # Map employment info
        current_job = None
        if employment_history:
            # Find current job (most recent or marked as present)
            for job in employment_history:
                if job.get("end_date") == "present" or not job.get("end_date"):
                    current_job = job
                    break
            
            if not current_job:
                current_job = employment_history[0]  # Most recent
        
        assessment_data["employment_info"].update({
            "current_status": data.get("current_employment_status", "unemployed"),
            "history": employment_history,
            "total_experience_years": data.get("total_experience_years", 0),
            "skills": data.get("skills", []),
            "education_level": self._determine_education_level(education)
        })
        
        # Update demographic info
        assessment_data["demographic_info"].update({
            "education_level": self._determine_education_level(education),
            "languages": data.get("languages", [])
        })
    
    async def _map_assets_data(self, data: Dict[str, Any], assessment_data: Dict[str, Any]):
        """Map assets/liabilities data to assessment format."""
        assets = data.get("assets", {})
        liabilities = data.get("liabilities", {})
        
        assessment_data["wealth_info"].update({
            "savings": assets.get("cash_and_equivalents", 0),
            "property_value": sum([prop.get("estimated_value", 0) for prop in assets.get("property", [])]),
            "investments": {
                "stocks": assets.get("investments", {}).get("stocks", 0),
                "bonds": assets.get("investments", {}).get("bonds", 0),
                "mutual_funds": assets.get("investments", {}).get("mutual_funds", 0),
                "real_estate": assets.get("investments", {}).get("real_estate", 0)
            },
            "total_assets": assets.get("total_assets", 0),
            "total_debts": liabilities.get("total_liabilities", 0),
            "net_worth": data.get("net_worth", 0),
            "credit_cards": liabilities.get("credit_cards", []),
            "loans": liabilities.get("loans", [])
        })
    
    async def _map_credit_data(self, data: Dict[str, Any], assessment_data: Dict[str, Any]):
        """Map credit report data to assessment format."""
        personal_info = data.get("personal_info", {})
        credit_score = data.get("credit_score", {})
        credit_utilization = data.get("credit_utilization", {})
        payment_history = data.get("payment_history", {})
        contact_info = data.get("contact_info", {})
        addresses = data.get("addresses", [])
        
        # Update applicant info if not already set
        if not assessment_data["applicant_info"].get("name") and personal_info.get("name"):
            assessment_data["applicant_info"]["name"] = personal_info.get("name", "")
        
        if not assessment_data["applicant_info"].get("age") and personal_info.get("date_of_birth"):
            assessment_data["applicant_info"]["age"] = self._calculate_age(personal_info.get("date_of_birth"))
        
        # Update demographic info
        assessment_data["demographic_info"].update({
            "gender": personal_info.get("gender", ""),
            "nationality": personal_info.get("nationality", ""),
            "emirates_id": personal_info.get("emirates_id", ""),
            "passport_number": personal_info.get("passport_number", "")
        })
        
        # Update contact info
        if contact_info.get("email"):
            assessment_data["applicant_info"]["email"] = contact_info.get("email")
        
        if contact_info.get("phone_numbers") or contact_info.get("mobile_numbers"):
            phones = contact_info.get("phone_numbers", []) + contact_info.get("mobile_numbers", [])
            if phones:
                assessment_data["applicant_info"]["phone"] = phones[0]  # Use first phone number
        
        # Update address if not already set
        if not assessment_data["applicant_info"].get("address") and addresses:
            latest_address = addresses[0]  # Assume first is most recent
            assessment_data["applicant_info"]["address"] = latest_address.get("address", "")
            if latest_address.get("emirate"):
                assessment_data["demographic_info"]["emirate"] = latest_address.get("emirate")
        
        # Update wealth info with credit data
        assessment_data["wealth_info"].update({
            "credit_score": credit_score.get("score", 0),
            "credit_rating": credit_score.get("rating", ""),
            "credit_bureau": credit_score.get("bureau", "AECB"),
            "credit_utilization_ratio": credit_utilization.get("utilization_ratio", 0),
            "payment_history": {
                "on_time_percentage": payment_history.get("on_time_payments", 0),
                "late_payments": payment_history.get("late_payments", 0),
                "defaults": payment_history.get("defaults", 0)
            }
        })
    
    async def _map_bank_statement_data(self, data: Dict[str, Any], assessment_data: Dict[str, Any]):
        """Map bank statement data to assessment format."""
        account_info = data.get("account_info", {})
        balance_summary = data.get("balance_summary", {})
        transaction_summary = data.get("transaction_summary", {})
        income_analysis = data.get("income_analysis", {})
        spending_analysis = data.get("spending_analysis", {})
        financial_behavior = data.get("financial_behavior", {})
        
        # Update applicant info if not already set
        if not assessment_data["applicant_info"].get("name") and account_info.get("account_holder_name"):
            assessment_data["applicant_info"]["name"] = account_info.get("account_holder_name", "")
        
        # Update income info
        assessment_data["income_info"].update({
            "monthly_income": income_analysis.get("average_monthly_salary", 0),
            "annual_income": income_analysis.get("average_monthly_salary", 0) * 12,
            "income_sources": ["salary"] + income_analysis.get("other_income_sources", []),
            "income_stability": "stable" if income_analysis.get("salary_frequency") == "monthly" else "irregular",
            "bank_account_balance": balance_summary.get("closing_balance", 0),
            "average_balance": balance_summary.get("average_balance", 0)
        })
        
        # Update wealth info with banking data
        assessment_data["wealth_info"].update({
            "bank_accounts": [{
                "bank_name": account_info.get("bank_name", ""),
                "account_type": account_info.get("account_type", ""),
                "balance": balance_summary.get("closing_balance", 0),
                "account_number": account_info.get("account_number", "")[-4:] if account_info.get("account_number") else ""  # Last 4 digits only
            }],
            "monthly_income": income_analysis.get("total_monthly_income", 0),
            "monthly_expenses": spending_analysis.get("average_monthly_spending", 0),
            "cash_flow": income_analysis.get("total_monthly_income", 0) - spending_analysis.get("average_monthly_spending", 0),
            "financial_behavior": {
                "account_management": financial_behavior.get("account_management", ""),
                "overdraft_incidents": financial_behavior.get("overdraft_incidents", 0),
                "bounce_incidents": financial_behavior.get("bounce_incidents", 0),
                "cash_flow_pattern": financial_behavior.get("cash_flow_pattern", ""),
                "atm_usage_frequency": spending_analysis.get("atm_withdrawal_frequency", 0)
            },
            "banking_relationship": {
                "primary_bank": account_info.get("bank_name", ""),
                "account_age_months": 12,  # Default, could be calculated from statement period
                "transaction_volume": transaction_summary.get("total_transactions", 0),
                "service_charges": spending_analysis.get("service_charges_monthly", 0)
            }
        })
        
        # Update family info if we can infer anything
        # Note: Bank statements don't typically contain family info directly
        # but we might infer some patterns from transaction descriptions
        
        # Update employment info if salary patterns are detected
        if income_analysis.get("average_monthly_salary", 0) > 0:
            assessment_data["employment_info"].update({
                "current_status": "employed",
                "monthly_salary": income_analysis.get("average_monthly_salary", 0),
                "salary_frequency": income_analysis.get("salary_frequency", "monthly"),
                "employment_stability": "stable" if financial_behavior.get("cash_flow_pattern") == "stable" else "unstable"
            })
    
    def _calculate_age(self, date_of_birth: str) -> int:
        """Calculate age from date of birth."""
        if not date_of_birth:
            return 0
        
        try:
            from datetime import datetime
            birth_date = datetime.strptime(date_of_birth, "%Y-%m-%d")
            today = datetime.now()
            age = today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
            return age
        except:
            return 0
    
    def _determine_education_level(self, education: List[Dict[str, Any]]) -> str:
        """Determine highest education level."""
        if not education:
            return "unknown"
        
        levels = {
            "phd": 6, "doctorate": 6, "doctoral": 6,
            "master": 5, "masters": 5, "mba": 5,
            "bachelor": 4, "bachelors": 4, "degree": 4,
            "diploma": 3, "associate": 3,
            "high school": 2, "secondary": 2,
            "primary": 1, "elementary": 1
        }
        
        highest_level = 0
        highest_degree = "unknown"
        
        for edu in education:
            degree = edu.get("degree", "").lower()
            for level_name, level_value in levels.items():
                if level_name in degree:
                    if level_value > highest_level:
                        highest_level = level_value
                        highest_degree = level_name
                    break
        
        return highest_degree if highest_degree != "unknown" else "high_school"
    
    def cleanup_temp_files(self):
        """Clean up temporary files."""
        try:
            for file_path in self.temp_dir.glob("temp_*"):
                file_path.unlink()
        except Exception as e:
            logger.warning(f"Failed to cleanup temp files: {str(e)}")


# Test implementation
if __name__ == "__main__":
    import asyncio
    
    async def test_document_processor():
        """Test the document processing agent."""
        print("🧪 Testing Document Processing Agent...")
        
        # Create agent
        agent = DocumentProcessingAgent(model="gpt-4o")
        
        # Test with sample files (using relative paths)
        test_files = [
            {
                "file_path": "sample_emirates_id.pdf",
                "document_purpose": "emirates_id"
            },
            {
                "file_path": "sample_resume.pdf", 
                "document_purpose": "resume"
            },
            {
                "file_path": "sample_assets.xlsx",
                "document_purpose": "assets_liabilities"
            },
            {
                "file_path": "sample_credit_report.txt",
                "document_purpose": "credit_report"
            },
            {
                "file_path": "sample_bank_statement.txt",
                "document_purpose": "bank_statement"
            }
        ]
        
        print('\nCurrent Working Directory -', os.getcwd())
        print('\nList of Directory -', os.listdir())
        
        # Check which files actually exist
        print('\n📁 Checking file availability:')
        for test_file in test_files:
            file_path = test_file["file_path"]
            exists = os.path.exists(file_path)
            print(f"  {file_path}: {'✅ Found' if exists else '❌ Not found'}")
            
            # Also check absolute path
            abs_path = os.path.abspath(file_path)
            print(f"    Absolute path: {abs_path}")
            
            # If file doesn't exist, try to find it in current directory
            if not exists:
                filename = os.path.basename(file_path)
                for root, dirs, files in os.walk('.'):
                    if filename in files:
                        found_path = os.path.join(root, filename)
                        print(f"    Found alternative: {found_path}")
                        test_file["file_path"] = found_path
                        break

        processed_docs = []
        
        for test_file in test_files:
            if os.path.exists(test_file["file_path"]):
                print(f"📄 Processing {test_file['file_path']}...")
                
                try:
                    result = await agent.execute(test_file)
                    processed_docs.append(result.get("extracted_data", {}))
                    
                    print(f"✅ Processed successfully!")
                    print(f"Document Type: {result.get('extracted_data', {}).get('document_type')}")
                    print(f"Confidence: {result.get('extracted_data', {}).get('confidence_score', 0):.2f}")
                    
                except Exception as e:
                    print(f"❌ Processing failed: {str(e)}")
            else:
                print(f"⚠️  File not found: {test_file['file_path']}")
        
        # Test conversion to assessment format
        if processed_docs:
            print(f"\n🔄 Converting to assessment format...")
            assessment_data = await agent.convert_to_assessment_format(processed_docs, "TEST-APP-001")
            
            print(f"✅ Conversion complete!")
            print(f"Application ID: {assessment_data.get('application_id')}")
            print(f"Applicant: {assessment_data.get('applicant_info', {}).get('name', 'N/A')}")
            print(f"Document Sources: {len(assessment_data.get('document_sources', []))}")
            print(f"Overall Response:", assessment_data)
        
        # Cleanup
        agent.cleanup_temp_files()
        print(f"\n🧹 Cleanup complete!")
    
    # Run test
    asyncio.run(test_document_processor())